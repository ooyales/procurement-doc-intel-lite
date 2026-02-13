import io
import csv
import logging

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Layer 1: Extract raw tables and text from uploaded documents."""

    def extract(self, file_path: str, file_format: str) -> dict:
        """Returns {'tables': [list of rows], 'full_text': str, 'page_count': int, 'method': str}"""
        fmt = file_format.lower()
        if fmt == 'pdf':
            return self._extract_pdf(file_path)
        elif fmt in ('xlsx', 'xls'):
            return self._extract_excel(file_path)
        elif fmt in ('docx', 'doc'):
            return self._extract_docx(file_path)
        elif fmt == 'csv':
            return self._extract_csv(file_path)
        else:
            raise ValueError(f"Unsupported format: {file_format}")

    def _extract_pdf(self, path):
        import pdfplumber
        tables = []
        full_text_parts = []
        page_count = 0
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text_parts.append(text)
                page_tables = page.extract_tables()
                if page_tables:
                    for t in page_tables:
                        # Convert to list of dicts using first row as headers
                        if len(t) >= 2:
                            headers = [str(h or '').strip() for h in t[0]]
                            for row in t[1:]:
                                row_dict = {}
                                for i, cell in enumerate(row):
                                    if i < len(headers) and headers[i]:
                                        row_dict[headers[i]] = str(cell or '').strip()
                                if any(row_dict.values()):
                                    tables.append(row_dict)
        return {
            'tables': tables,
            'full_text': '\n\n'.join(full_text_parts),
            'page_count': page_count,
            'method': 'pdfplumber',
        }

    def _extract_excel(self, path):
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        tables = []
        full_text_parts = []
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            # Find header row (first row with multiple non-empty cells)
            header_idx = 0
            for i, row in enumerate(rows):
                non_empty = sum(1 for c in row if c is not None and str(c).strip())
                if non_empty >= 2:
                    header_idx = i
                    break
            headers = [str(h or '').strip() for h in rows[header_idx]]
            for row in rows[header_idx + 1:]:
                row_dict = {}
                row_text_parts = []
                for j, cell in enumerate(row):
                    val = str(cell or '').strip()
                    if j < len(headers) and headers[j]:
                        row_dict[headers[j]] = val
                    if val:
                        row_text_parts.append(val)
                if any(row_dict.values()):
                    tables.append(row_dict)
                if row_text_parts:
                    full_text_parts.append(' | '.join(row_text_parts))
        return {
            'tables': tables,
            'full_text': '\n'.join(full_text_parts),
            'page_count': len(wb.worksheets),
            'method': 'openpyxl',
        }

    def _extract_docx(self, path):
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        tables = []
        full_text_parts = []
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text_parts.append(text)
        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if len(rows) >= 2:
                headers = rows[0]
                for row in rows[1:]:
                    row_dict = {}
                    for i, cell in enumerate(row):
                        if i < len(headers) and headers[i]:
                            row_dict[headers[i]] = cell
                    if any(row_dict.values()):
                        tables.append(row_dict)
        return {
            'tables': tables,
            'full_text': '\n\n'.join(full_text_parts),
            'page_count': 1,
            'method': 'python-docx',
        }

    def _extract_csv(self, path):
        import pandas as pd
        df = pd.read_csv(path)
        tables = df.fillna('').to_dict('records')
        full_text = df.to_string(index=False)
        return {
            'tables': tables,
            'full_text': full_text,
            'page_count': 1,
            'method': 'pandas',
        }
